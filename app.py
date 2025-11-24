# app.py
import os
import asyncio
import io
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from google.adk.models import Gemini
from google.adk.agents import Agent
from google.adk.tools.agent_tool import AgentTool
from google.adk.tools import google_search
from google.adk.runners import InMemoryRunner

# -----------------------------------------------------------------------------
# 0. CONFIG: API key + retry config
# -----------------------------------------------------------------------------

# Check if API key is already set
api_key_from_env = "GOOGLE_API_KEY" in os.environ and os.environ["GOOGLE_API_KEY"]

# Sidebar for API key
with st.sidebar:
    st.header("⚙️ Configuration")
    
    if api_key_from_env:
        st.success("✅ API Key: Set from environment variable")
        st.caption("Your Gemini API key is configured.")
    else:
        st.warning("⚠️ API Key: Not set")
        api_key_input = st.text_input(
            "Enter Google API Key",
            type="password",
            help="Get your API key from https://aistudio.google.com/app/apikey",
        )
        if api_key_input:
            os.environ["GOOGLE_API_KEY"] = api_key_input
            st.success("✅ API Key saved for this session")
        else:
            st.error("Please enter your API key to use the app")
    
    st.markdown("---")
    st.markdown("### 📚 About")
    st.markdown("This app uses AI to help with job applications by generating tailored CVs, cover letters, and interview prep.")


# Note: retry_options is not supported by the Gemini model class
# The model will use default retry behavior


# -----------------------------------------------------------------------------
# 1. HELPER: Convert Markdown to Word Document
# -----------------------------------------------------------------------------

def markdown_to_docx(markdown_text: str, title: str = "") -> bytes:
    """Convert markdown text to a Word document and return as bytes."""
    doc = Document()
    
    # Add title if provided
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Split content by lines and process
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Handle headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Handle numbered lists
        elif line[0:2].replace('.', '').isdigit():
            doc.add_paragraph(line, style='List Number')
        # Regular paragraph
        else:
            # Remove markdown bold/italic markers for cleaner output
            clean_line = line.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            doc.add_paragraph(clean_line)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


# -----------------------------------------------------------------------------
# 2. HELPERS: event → state extraction + async runner
# -----------------------------------------------------------------------------

def extract_state_from_events(events):
    """Extract state from ADK events, with fallback to event content."""
    final_state = {}
    last_content = None
    
    if not isinstance(events, (list, tuple)):
        events = [events]
    
    for e in events:
        # Try to get state_delta
        actions = getattr(e, "actions", None)
        if actions and getattr(actions, "state_delta", None):
            final_state.update(actions.state_delta)
        
        # Also capture the last event's content as fallback
        if hasattr(e, "content"):
            last_content = e.content
        elif hasattr(e, "text"):
            last_content = e.text
        elif hasattr(e, "message"):
            msg = e.message
            if hasattr(msg, "content"):
                last_content = msg.content
            elif hasattr(msg, "text"):
                last_content = msg.text
    
    # If we didn't get any state, use the last content
    if not final_state and last_content:
        final_state = {"output": last_content}
    
    return final_state


async def run_agent(agent, prompt: str):
    """Run a single ADK agent and return its state dict."""
    runner = InMemoryRunner(agent=agent)
    
    # Use run_debug to get events
    events = await runner.run_debug(prompt, quiet=True, verbose=False)
    
    # Extract state from events
    state = extract_state_from_events(events)
    
    # If the agent has an output_key, try to extract it
    if hasattr(agent, "output_key") and agent.output_key:
        # If output_key is not in state, check if we have generic 'output'
        if agent.output_key not in state and "output" in state:
            state[agent.output_key] = state["output"]
    
    # Debug: show what we extracted
    if not state or all(not v for v in state.values()):
        st.warning(f"⚠️ Agent '{agent.name}' returned empty state. Check debug output.")
    
    return state


# -----------------------------------------------------------------------------
# 2. AGENT DEFINITIONS
#    (Paste or adapt from your notebook; these are example versions)
# -----------------------------------------------------------------------------

job_analyst_agent = Agent(
    name="job_analyst_agent",
    model=Gemini(model="gemini-2.0-flash-exp"),
    instruction="""
You are an ATS-style job analysis agent.

You will receive a single message that contains BOTH:
- A job description
- A candidate CV

Your tasks:
1. Identify and list:
   - Hard Skills & Technologies
   - Soft Skills
   - Key Qualifications & Education
2. Extract the Top 15 ATS Keywords (comma-separated).
3. Estimate an ATS match score (0–100%).

Output format (markdown):

### Hard Skills & Technologies
- ...

### Soft Skills
- ...

### Key Qualifications & Education
- ...

### Top 15 ATS Keywords
keyword1, keyword2, ...

### ATS Match Score
XX%

Return ONLY this analysis in markdown.
""",
    output_key="keywords",
)

cvwriter_agent = Agent(
    name="cvwriter_agent",
    model=Gemini(model="gemini-2.0-flash-exp"),
    instruction="""
You are a CV rewriting agent.

You will receive:
- A job description
- A current CV

Task:
- Rewrite the CV so it is clearly tailored to this role.
- Use metrics where possible.
- Keep the format clean and ATS-friendly.
- DO NOT repeat the full job description.

Return ONLY the revised CV.
""",
    output_key="revisedcv",
)

clwriter_agent = Agent(
    name="clwriter_agent",
    model=Gemini(model="gemini-2.0-flash-exp"),
    instruction="""
You are a cover letter writing agent.

You will receive:
- A job description
- A current CV

Task:
- Write a professional cover letter for this specific role and candidate.
- Use a clear business letter format.
- Base it on the job ad and the candidate profile.
- Do NOT repeat the full job description or CV.

Return ONLY the cover letter.
""",
    output_key="coverletter",
)

coach_agent = Agent(
    name="coach_agent",
    model=Gemini(model="gemini-2.0-flash-exp"),
    instruction="""
You are an interview coach.

You will receive:
- A job description
- A candidate CV

Produce an interview preparation guide with THREE sections:

## Behavioural / Fit Questions
- 5–7 questions
- For each, a strong sample answer in the first person ("I ...")

## Technical / Analytics Questions
- 5–7 questions on SQL, dashboards, CX metrics, A/B testing, survey analysis, etc.
- For each, a strong sample answer tailored to the candidate

## Questions to Ask the Employer
- 5 thoughtful questions the candidate can ask

Do NOT repeat the full job ad or CV.
Return ONLY this guide in markdown.
""",
    output_key="interview_guide",
)

research_agent = Agent(
    name="research_agent",
    model=Gemini(model="gemini-2.0-flash-exp"),
    instruction="""
You are a job search agent.

You MUST use the google_search tool.

You will receive a job description. Find 3 real job postings that are similar:
- Titles: Data & Insights Analyst, Data Analyst, CX Analyst, Customer Insights Analyst, etc.
- Location: Toronto, ON OR Remote roles open to candidates in Canada.

For EACH job, return:
- Job Title – Company (Location) — URL

Output ONLY a markdown bullet list. No extra explanation.
""",
    tools=[google_search],
    output_key="similar_jobs",
)


# -----------------------------------------------------------------------------
# 3. STREAMLIT UI
# -----------------------------------------------------------------------------

st.set_page_config(page_title="Job Application Assistant", layout="wide")

# Display logo and title
col_logo, col_title = st.columns([1, 4])

with col_logo:
    st.image("logo.jpg", width=150)

with col_title:
    st.title("🤖 Job Application Assistant")
    st.write("Paste your job ad and CV, choose what you want, and I'll generate tailored outputs.")

st.markdown("---")

col1, col2 = st.columns(2)

with col1:
    job_ad_text = st.text_area(
        "Job Description / Job Ad",
        height=350,
        placeholder="Paste the full job ad here...",
    )

with col2:
    base_cv_text = st.text_area(
        "Current CV / Resume",
        height=350,
        placeholder="Paste your current CV here...",
    )

st.markdown("---")

tasks = st.multiselect(
    "What do you want me to generate?",
    options=[
        "ATS analysis",
        "Revised CV",
        "Cover letter",
        "Interview guide",
        "Similar jobs",
    ],
    default=["ATS analysis", "Revised CV", "Cover letter"],
    help="Pick one or more tasks.",
)

run_button = st.button("🚀 Run Assistant")


# -----------------------------------------------------------------------------
# 4. ORCHESTRATION: run selected agents
# -----------------------------------------------------------------------------

async def run_all(job_ad: str, cv_text: str, tasks_selected):
    context_prompt = f"""
========= JOB DESCRIPTION =========
{job_ad}

========= CURRENT CV =========
{cv_text}
"""

    results = {}

    # Always do ATS analysis if chosen
    if "ATS analysis" in tasks_selected:
        state = await run_agent(job_analyst_agent, context_prompt)
        results["keywords"] = state.get("keywords", "No ATS analysis found.")

    if "Revised CV" in tasks_selected:
        cv_prompt = f"""
You will receive a job description and a current CV.

Your task:
- Rewrite the CV so it is clearly tailored to this role.
- Use the job description to emphasise relevant skills and experience.
- Use metrics where possible.
- Keep the format clean and ATS-friendly.

Return ONLY the revised CV.

{context_prompt}
"""
        state = await run_agent(cvwriter_agent, cv_prompt)
        results["revisedcv"] = state.get("revisedcv", "No revised CV found.")

    if "Cover letter" in tasks_selected:
        cl_prompt = f"""
You will receive a job description and a current CV.

Your task:
- Write a professional cover letter for this specific role and candidate.
- Use a clear business letter format.
- Base the content on the job description and the candidate's profile.

Return ONLY the final cover letter.

{context_prompt}
"""
        state = await run_agent(clwriter_agent, cl_prompt)
        results["coverletter"] = state.get("coverletter", "No cover letter found.")

    if "Interview guide" in tasks_selected:
        coach_prompt = f"""
You will receive a job description and a current CV.

Your task:
- Generate an interview preparation guide for this role and candidate:
  - 5–7 Behavioural / Fit questions + sample answers
  - 5–7 Technical / Analytics questions + sample answers
  - 5 smart questions the candidate should ask the employer

Return ONLY the interview preparation guide in markdown.

{context_prompt}
"""
        state = await run_agent(coach_agent, coach_prompt)
        results["interview_guide"] = state.get("interview_guide", "No interview guide found.")

    if "Similar jobs" in tasks_selected:
        jobs_prompt = f"""
Here is a job description:

{job_ad}

Use google_search to find 3 similar jobs:
- Similar title (Data & Insights Analyst, CX Analyst, Customer Insights Analyst, etc.)
- Located in Toronto, ON or Remote (Canada)
Return ONLY a markdown bullet list:
- Job Title – Company (Location) — URL
"""
        state = await run_agent(research_agent, jobs_prompt)
        results["similar_jobs"] = state.get("similar_jobs", "No similar jobs found.")

    return results


if run_button:
    if not job_ad_text or not base_cv_text:
        st.error("Please paste both the job ad and your CV before running.")
    elif not tasks:
        st.error("Please select at least one task to run.")
    else:
        try:
            with st.spinner("Running agents..."):
                outputs = asyncio.run(run_all(job_ad_text, base_cv_text, tasks))

            st.markdown("## ✅ Results")
            
            # Enhanced debug expander to show raw outputs
            with st.expander("🔍 Debug: View Raw Agent Outputs"):
                st.markdown("### Full Outputs Dictionary")
                st.json(outputs)
                
                st.markdown("### Individual Output Details")
                for key, value in outputs.items():
                    st.markdown(f"**{key}**:")
                    st.write(f"- Type: `{type(value).__name__}`")
                    st.write(f"- Length: `{len(str(value)) if value else 0}` characters")
                    st.write(f"- Is Empty: `{not value or value.startswith('No ')}`")
                    if value and not value.startswith("No "):
                        st.text_area(f"Content preview ({key})", value[:500] + "..." if len(value) > 500 else value, height=100, key=f"debug_{key}")

            if "ATS analysis" in tasks:
                st.markdown("### ATS / Keywords")
                keywords_content = outputs.get("keywords", "No ATS analysis found.")
                if not keywords_content or keywords_content == "No ATS analysis found.":
                    st.warning("⚠️ No ATS analysis was generated. The agent may not have returned content.")
                st.markdown(keywords_content, unsafe_allow_html=False)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📥 Download as Markdown",
                        data=keywords_content if keywords_content else "No content available",
                        file_name="ats_analysis.md",
                        mime="text/markdown",
                        key="download_ats_md"
                    )
                with col2:
                    if keywords_content and keywords_content != "No ATS analysis found.":
                        docx_data = markdown_to_docx(keywords_content, "ATS Analysis")
                        st.download_button(
                            label="📄 Download as Word",
                            data=docx_data,
                            file_name="ats_analysis.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_ats_docx"
                        )

            if "Revised CV" in tasks:
                st.markdown("### Revised CV")
                cv_content = outputs.get("revisedcv", "No revised CV found.")
                if not cv_content or cv_content == "No revised CV found.":
                    st.warning("⚠️ No revised CV was generated. The agent may not have returned content.")
                st.markdown(cv_content)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📥 Download as Markdown",
                        data=cv_content if cv_content else "No content available",
                        file_name="revised_cv.md",
                        mime="text/markdown",
                        key="download_cv_md"
                    )
                with col2:
                    if cv_content and cv_content != "No revised CV found.":
                        docx_data = markdown_to_docx(cv_content, "Curriculum Vitae")
                        st.download_button(
                            label="📄 Download as Word",
                            data=docx_data,
                            file_name="revised_cv.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_cv_docx"
                        )

            if "Cover letter" in tasks:
                st.markdown("### Cover Letter")
                cl_content = outputs.get("coverletter", "No cover letter found.")
                if not cl_content or cl_content == "No cover letter found.":
                    st.warning("⚠️ No cover letter was generated. The agent may not have returned content.")
                st.markdown(cl_content)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📥 Download as Markdown",
                        data=cl_content if cl_content else "No content available",
                        file_name="cover_letter.md",
                        mime="text/markdown",
                        key="download_cl_md"
                    )
                with col2:
                    if cl_content and cl_content != "No cover letter found.":
                        docx_data = markdown_to_docx(cl_content, "Cover Letter")
                        st.download_button(
                            label="📄 Download as Word",
                            data=docx_data,
                            file_name="cover_letter.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_cl_docx"
                        )

            if "Interview guide" in tasks:
                st.markdown("### Interview Preparation Guide")
                guide_content = outputs.get("interview_guide", "No interview guide found.")
                if not guide_content or guide_content == "No interview guide found.":
                    st.warning("⚠️ No interview guide was generated. The agent may not have returned content.")
                st.markdown(guide_content)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📥 Download as Markdown",
                        data=guide_content if guide_content else "No content available",
                        file_name="interview_guide.md",
                        mime="text/markdown",
                        key="download_guide_md"
                    )
                with col2:
                    if guide_content and guide_content != "No interview guide found.":
                        docx_data = markdown_to_docx(guide_content, "Interview Preparation Guide")
                        st.download_button(
                            label="📄 Download as Word",
                            data=docx_data,
                            file_name="interview_guide.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_guide_docx"
                        )

            if "Similar jobs" in tasks:
                st.markdown("### Similar Jobs")
                jobs_content = outputs.get("similar_jobs", "No similar jobs found.")
                if not jobs_content or jobs_content == "No similar jobs found.":
                    st.warning("⚠️ No similar jobs were found. The agent may not have returned content.")
                st.markdown(jobs_content)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📥 Download as Markdown",
                        data=jobs_content if jobs_content else "No content available",
                        file_name="similar_jobs.md",
                        mime="text/markdown",
                        key="download_jobs_md"
                    )
                with col2:
                    if jobs_content and jobs_content != "No similar jobs found.":
                        docx_data = markdown_to_docx(jobs_content, "Similar Job Opportunities")
                        st.download_button(
                            label="📄 Download as Word",
                            data=docx_data,
                            file_name="similar_jobs.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_jobs_docx"
                        )
        
        except Exception as e:
            st.error(f"An error occurred while running the agents: {str(e)}")
            st.exception(e)


